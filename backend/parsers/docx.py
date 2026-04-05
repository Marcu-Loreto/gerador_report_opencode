from .base import BaseParser, ParsedDocument
import io


class DOCXParser(BaseParser):
    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        try:
            from docx import Document

            doc = Document(io.BytesIO(content))
            text = "\n".join([para.text for para in doc.paragraphs])

            tables_text = ""
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    tables_text += " | ".join(cells) + "\n"

            full_text = text + "\n" + tables_text

            return ParsedDocument(
                texto=full_text,
                estrutura=self._detect_structure(text),
                metadados={
                    "arquivo": filename,
                    "tamanho": len(content),
                    "paragrafos": len(doc.paragraphs),
                    "tabelas": len(doc.tables),
                },
                alertas=[],
                limitacoes=[],
            )
        except Exception as e:
            result = self._create_empty_result(filename)
            result.alertas.append(f"Erro ao processar DOCX: {str(e)}")
            return result
