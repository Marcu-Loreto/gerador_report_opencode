from typing import Optional
from io import BytesIO
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_to_docx(content: str, filename: str = "relatorio.docx") -> tuple[bytes, str]:
    try:
        doc = Document()

        lines = content.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith("# "):
                heading = doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                heading = doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                heading = doc.add_heading(line[4:], level=3)
            else:
                p = doc.add_paragraph(line)

        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)

        return doc_stream.getvalue(), f"{Path(filename).stem}.docx"
    except Exception as e:
        raise RuntimeError(f"Erro ao gerar DOCX: {str(e)}")


def export_to_docx_with_style(
    content: str, filename: str = "relatorio.docx"
) -> tuple[bytes, str]:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    title = doc.add_heading("Relatório", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = content.split("\n")
    in_list = False

    for line in lines:
        line = line.strip()
        if not line:
            in_list = False
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            if not in_list:
                in_list = True
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            p = doc.add_paragraph(line)
            in_list = False

    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    return doc_stream.getvalue(), f"{Path(filename).stem}.docx"
